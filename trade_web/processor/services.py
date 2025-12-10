from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import requests
from bs4 import BeautifulSoup
import io
import re

def extract_abbreviations(doc):
    """
    Extract СОКРАЩЕНИЯ И СИМВОЛЫ from Table 1
    Returns: dict {abbreviation: full_name}
    """
    abbreviations = {}
    if len(doc.tables) < 2:
        return abbreviations
        
    table = doc.tables[1]

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 2 and cells[0] and cells[1]:
            abbr = cells[0].replace('\n', ' / ')
            full_name = cells[1].replace('\n', ' / ')
            abbreviations[abbr] = full_name

    return abbreviations

def extract_units(doc):
    """
    Extract ЕДИНИЦЫ ИЗМЕРЕНИЯ from Table 2
    Returns: list of dicts with keys: name, symbol, code
    """
    units = []
    if len(doc.tables) < 3:
        return units
        
    table = doc.tables[2]

    for row in table.rows[1:]:  # Skip header
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 3:
            name, symbol, code = cells[0], cells[1], cells[2]

            # Skip section headers and empty rows
            if name == symbol == code or not all([name, symbol, code]):
                continue

            units.append({
                "name": name,
                "symbol": symbol,
                "code": code
            })

    return units

def extract_tn_ved_codes(doc):
    """
    Extract all TN VED codes from tables with header containing 'Код ТН ВЭД'
    Returns: list of dicts with keys: code, name, unit
    """
    tn_ved_codes = []

    # Find all TN VED tables
    for table in doc.tables:
        if not table.rows:
            continue

        header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if 'Код' not in header_text or 'ТН ВЭД' not in header_text:
            continue

        # Extract data from this table
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) >= 4:
                code = cells[1]
                name = cells[2]
                unit = cells[3]
            elif len(cells) >= 3:
                code = cells[0] if cells[0] else cells[1]
                name = cells[1] if cells[0] else cells[2]
                unit = cells[2] if cells[0] else ''
            else:
                continue

            # Clean up
            code = code.replace('\xa0', '').replace(' ', '')
            name = name.replace('\xa0', ' ').strip()
            unit = unit.replace('\xa0', ' ').strip()

            if not name:
                continue

            tn_ved_codes.append({
                "code": code,
                "name": name,
                "unit": unit
            })

    return tn_ved_codes

def extract_all(docx_file):
    """
    Extract all data from the document
    Returns: tuple (abbreviations, units, tn_ved_codes)
    """
    doc = Document(docx_file)

    abbreviations = extract_abbreviations(doc)
    units = extract_units(doc)
    tn_ved_codes = extract_tn_ved_codes(doc)

    return abbreviations, units, tn_ved_codes

def translate_rate(text):
    """
    Translate duty rate expressions from Uzbek to Russian, covering various complex formats.

    Args:
        text (str): The Uzbek duty rate expression.

    Returns:
        str: The translated Russian duty rate expression.
    """
    if not text:
        return ""

    # 1. Handle simple single-value/marked cases (e.g., "10*", "20***")
    # This should be done before complex regex to avoid false matches.
    if re.fullmatch(r'[\d,\s*]+[\*]+', text.strip()):
        return text.strip()

    # 2. Basic currency and unit replacements
    # Prioritize 'АҚШ доллари' (full form) over 'АҚШ долл.' (abbreviated)
    text = text.replace("АҚШ доллари", "долл. США")
    text = text.replace("АҚШ долл.", "долл. США")
    text = text.replace("куб. см.", "куб. см.")  # Keeps it the same, just for completeness of units
    text = text.replace("/кг", " за кг")

    # 3. Complex phrase: "X, лекин ... кам бўлмаган миқдорда"
    # Matches: '..., лекин <description> <amount> АҚШ долларидан кам бўлмаган миқдорда'
    # Captures:
    #   Group 1: The first part (e.g., "20")
    #   Group 2: The descriptive part (e.g., "ҳар бир килограмми учун 0,3")
    match_min_amount = re.search(
        r'(.+),\s*лекин\s*(ҳар бир|ҳар)\s*(.*?)\s*дан кам бўлмаган миқдорда',
        text,
        re.IGNORECASE | re.DOTALL
    )
    if match_min_amount:
        part1 = match_min_amount.group(1).strip()  # e.g., "20"

        # Capture the part that contains the unit and amount (e.g., "килограмми учун 0,3 долл. США")
        # The capture group is already partially processed by the basic replacements.
        description_part = match_min_amount.group(2) + " " + match_min_amount.group(3)

        # Translate the descriptive unit part
        # Examples: "ҳар бир килограмми учун 0,3 долл. США" -> "не менее 0,3 долл. США за килограмм"
        # The key is to rearrange/translate the unit ('учун' means 'for')

        # Map Uzbek units/phrases to Russian
        unit_map = {
            r'ҳар бир килограмми учун ([\d\s.,]+)\s*(долл\. США)': r'но не менее \1 долл. США за килограмм',
            r'ҳар бир донаси учун ([\d\s.,]+)\s*(долл\. США)': r'но не менее \1 долл. США за штуку',
            r'ҳар бир литри учун ([\d\s.,]+)\s*(долл\. США)': r'но не менее \1 долл. США за литр',
            r'ҳар бир жуфти учун ([\d\s.,]+)\s*(долл\. США)': r'но не менее \1 долл. США за пару',
            r'ҳар бир м2 учун ([\d\s.,]+)\s*(долл\. США)': r'но не менее \1 долл. США за м2',
            r'ҳар 1000 донаси учун ([\d\s.,]+)\s*(долл\. США)': r'но не менее \1 долл. США за 1000 штук',
        }

        translated_description = description_part
        found_unit_match = False
        for uzb_pattern, rus_replacement in unit_map.items():
            # Use the unit_map to reformat and translate the descriptive part
            match = re.search(uzb_pattern, description_part)
            if match:
                # Replace the entire descriptive phrase with the Russian equivalent
                # Need to extract the amount and currency to insert into the Russian phrase
                amount = match.group(1).strip()
                # Create a specific pattern that captures everything before 'дан кам бўлмаган миқдорда'
                # and everything after 'лекин' but before the amount and currency.

                # Simple replacement for the most common structure:
                translated_description = rus_replacement.replace(r'\1', amount)
                found_unit_match = True
                break

        # If no specific unit match was found, fall back to a generic translation
        if not found_unit_match:
            # Handle the example: "15, лекин 0,15 АҚШ долл./кг дан кам эмас" (original structure)
            match_generic_min = re.search(r'(.*),\s*лекин\s*(.*)\s*дан кам эмас', text, re.IGNORECASE | re.DOTALL)
            if match_generic_min:
                part1 = match_generic_min.group(1).strip()
                part2 = match_generic_min.group(2).strip()
                return f"{part1}, но не менее {part2}"

            # For the complex structure that failed unit matching, just translate the main phrase structure
            # This is a fallback and might not be perfectly grammatical.
            translated_description = translated_description.replace("ҳар бир", "за").replace("учун", "").strip()
            return f"{part1}, но не менее {translated_description}"

        return f"{part1}, {translated_description}"

    # 4. Complex phrase: "X + Y АҚШ доллари ҳар бир куб. см. учун"
    # Matches: '... + ... АҚШ доллари ҳар бир куб. см. учун'
    # Captures:
    #   Group 1: The first rate (e.g., "70")
    #   Group 2: The additional rate (e.g., "3")
    #   Group 3: The unit marker (e.g., "куб. см. учун")
    match_additive = re.search(
        r'([\d\s.,]+)\s*\+\s*([\d\s.,]+)\s*(долл\. США)\s*ҳар бир\s*(.+)\s*учун([\*]*)',
        text,
        re.IGNORECASE | re.DOTALL
    )
    if match_additive:
        part1 = match_additive.group(1).strip()
        part2 = match_additive.group(2).strip()
        unit_part = match_additive.group(4).strip()
        stars = match_additive.group(5)

        # Translate unit (куб. см. учун -> за куб. см.)
        # Note: The 'ҳар бир' and 'учун' are translated to 'за' in Russian, and the order is flipped.
        if "куб. см." in unit_part:
            unit_part = "за куб. см."

        return f"{part1} + {part2} долл. США {unit_part}{stars}"

    # 5. Fallback for any other basic translation
    # If the text has no complex structure, return it after basic currency replacement
    return text.strip()

def fetch_duty_rates(url):
    import requests
    import urllib3
    from bs4 import BeautifulSoup
    import time

    print(f"Fetching duty rates from {url}...")

    # Disable warnings for verify=False
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Connection": "keep-alive",
        "Referer": "https://lex.uz/",
    }

    session = requests.Session()

    retries = 5
    html_content = None

    for attempt in range(1, retries + 1):
        try:
            print(f"Attempt {attempt}/{retries} ...")

            response = session.get(
                url,
                headers=headers,
                timeout=(5, 10),  # 5s to connect, 10s to read
                allow_redirects=True,
                verify=False
            )

            if response.status_code == 200:
                html_content = response.text
                print("Successfully fetched!")
                break

            print(f"Status: {response.status_code}, retrying...")

        except requests.exceptions.Timeout:
            print("Timeout, retrying...")
        except requests.exceptions.ConnectionError:
            print("Connection error, retrying...")
        except Exception as e:
            print(f"Unknown error: {e}, retrying...")

        time.sleep(1)

    if not html_content:
        print("Failed to fetch data from lex.uz")
        return {}


    soup = BeautifulSoup(html_content, 'html.parser')
    tables = soup.find_all('table')

    rates_map = {}

    for table in tables:
        headers = [th.get_text(strip=True) for th in table.find_all(['th', 'td'])]
        header_text = " ".join(headers)

        if (
            "ТИФ ТНнинг 2022 йилги таҳрири" in header_text
            and "Импорт божхона божи ставкаси" in header_text
        ):
            rows = table.find_all('tr')
            for row in rows:
                cells = row.find_all('td')
                if not cells:
                    continue

                cell_texts = [cell.get_text(strip=True) for cell in cells]

                if len(cell_texts) >= 3:
                    code_str = cell_texts[0]
                    rate = cell_texts[2]
                elif len(cell_texts) == 2:
                    code_str = cell_texts[0]
                    rate = cell_texts[1]
                else:
                    continue

                if "ТИФ ТНнинг" in code_str:
                    continue
                if len(rate) > 5:
                    print(rate)

                # Translate rate
                rate = translate_rate(rate)

                codes = [c.strip().replace(" ", "") for c in code_str.split(',')]
                for code in codes:
                    rates_map[code] = rate

    return rates_map

def get_duty_rate(code, rates_map):
    """
    Find duty rate for a code using longest prefix matching.
    """
    if not code:
        return ""
        
    # Clean code for matching
    clean_code = code.replace(' ', '')
    
    # Try to find match from longest to shortest
    for i in range(len(clean_code), 1, -1):
        prefix = clean_code[:i]
        if prefix in rates_map:
            return rates_map[prefix]
            
    return ""

def generate_tn_ved_excel(
        units,
        tn_ved_codes,
        duty_rates_url
):
    # Fetch duty rates
    rates_map = fetch_duty_rates(duty_rates_url)

    # Create unit → code map
    unit_code_map = {u['symbol']: u['code'] for u in units}
    unit_code_map[''] = ''  # Handle empty fields

    # Create workbook
    wb = Workbook()
    sheet = wb.active
    sheet.title = "ТН ВЭД"

    # Define styles
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="4472C4")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Set headers
    headers = ["Код ТН ВЭД", "Наименование позиции", "Доп. ед. изм.", "Код ед. изм.", "Ставка пошлины"]

    for col, header in enumerate(headers, 1):
        cell = sheet.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Fill rows
    for row_idx, entry in enumerate(tn_ved_codes, start=2):
        code = entry.get('code', '')
        name = entry.get('name', '')
        unit = entry.get('unit', '')
        unit_code = unit_code_map.get(unit, '')
        
        duty_rate = get_duty_rate(code, rates_map)
        
        sheet.cell(row=row_idx, column=1, value=code)
        sheet.cell(row=row_idx, column=2, value=name)
        sheet.cell(row=row_idx, column=3, value=unit)
        sheet.cell(row=row_idx, column=4, value=unit_code)
        sheet.cell(row=row_idx, column=5, value=duty_rate)

    # Column widths
    sheet.column_dimensions['A'].width = 20
    sheet.column_dimensions['B'].width = 70
    sheet.column_dimensions['C'].width = 15
    sheet.column_dimensions['D'].width = 12
    sheet.column_dimensions['E'].width = 15

    # Freeze header
    sheet.freeze_panes = "A2"

    return wb
